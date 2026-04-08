"""
Flask web dashboard for the job application pipeline.

Provides a modern UI to browse jobs, select for cover letter generation,
review generated content, and bulk-submit applications.
"""

import json
import os
import sys
import threading
from pathlib import Path

from flask import Flask, jsonify, render_template, request, Response, send_file

# Ensure project root is importable
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from utils.db import (
    init_db, get_job_by_id, get_jobs_by_status,
    get_queued_without_cl, get_queued_with_cl,
    update_job, count_by_status, get_contract_jobs,
)
from utils.profile import (
    load_profile, save_profile, validate_profile,
    load_preferences, save_preferences,
    get_default_profile, get_default_preferences,
    PROFILE_PATH, PREFERENCES_PATH,
)

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Pages
# ---------------------------------------------------------------------------

@app.route("/")
def dashboard():
    """Main dashboard page."""
    counts = count_by_status()
    return render_template("dashboard.html", counts=counts)


@app.route("/apply/<job_id>")
def apply_page(job_id):
    """Per-job application helper page with all fields ready to copy."""
    job = get_job_by_id(job_id)
    if not job:
        return "Job not found", 404

    profile = load_profile()
    name_parts = profile["name"].split(" ", 1)

    notes = {}
    if job.get("notes"):
        try:
            notes = json.loads(job["notes"])
        except (json.JSONDecodeError, TypeError):
            pass

    bullets = []
    if job.get("resume_bullets"):
        try:
            bullets = json.loads(job["resume_bullets"])
        except (json.JSONDecodeError, TypeError):
            pass

    return render_template("apply.html",
        job=job,
        profile=profile,
        first_name=name_parts[0],
        last_name=name_parts[1] if len(name_parts) > 1 else "",
        strengths=notes.get("strengths", []),
        missing=notes.get("missing", []),
        bullets=bullets,
    )


@app.route("/setup")
def setup_page():
    """Profile and preferences setup page."""
    return render_template("setup.html")


# ---------------------------------------------------------------------------
# API — Profile & Preferences
# ---------------------------------------------------------------------------

@app.route("/api/profile/full")
def api_profile_full():
    """Return the full candidate profile and preferences for the setup form.
    Falls back to baked-in defaults so a fresh clone is pre-filled."""
    profile = None
    try:
        profile = load_profile()
    except (FileNotFoundError, ValueError):
        pass

    try:
        preferences = load_preferences()
    except Exception:
        preferences = get_default_preferences()

    return jsonify({
        "exists": profile is not None,
        "profile": profile if profile else get_default_profile(),
        "preferences": preferences,
    })


@app.route("/api/resume/parse", methods=["POST"])
def api_resume_parse():
    """Upload a resume (PDF/DOCX/TXT), extract text, parse with AI into structured profile JSON."""
    import os
    import tempfile

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = (file.filename or "").lower()

    if not any(filename.endswith(ext) for ext in (".pdf", ".docx", ".doc", ".txt")):
        return jsonify({"error": "Supported formats: PDF, DOCX, TXT"}), 400

    # Save to temp file
    suffix = "." + filename.rsplit(".", 1)[-1]
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        # Extract text based on file type
        resume_text = _extract_resume_text(tmp_path, suffix)

        if not resume_text or len(resume_text.strip()) < 50:
            return jsonify({"error": "Could not extract enough text from the file"}), 400

        # Parse with AI
        profile = _parse_resume_with_ai(resume_text)
        profile["raw_resume_text"] = resume_text

        return jsonify({"profile": profile, "raw_text": resume_text})

    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        os.unlink(tmp_path)


def _extract_resume_text(file_path: str, suffix: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT file."""
    if suffix == ".pdf":
        from utils.resume_parser import parse_resume_pdf
        return parse_resume_pdf(file_path)

    elif suffix in (".docx", ".doc"):
        from docx import Document
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    else:  # .txt
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def _parse_resume_with_ai(resume_text: str) -> dict:
    """Send resume text to AI and get structured profile JSON back."""
    from utils.ai_client import chat_completion
    import re

    system = """You are a resume parser. Extract structured data from the resume text below.
Return valid JSON only with exactly these fields. If a field cannot be determined, use empty string or empty array.
Do NOT invent information that isn't in the resume."""

    user_msg = f"""Parse this resume into structured JSON:

{resume_text[:6000]}

Return this exact JSON structure:
{{
  "name": "Full Name",
  "email": "email@example.com",
  "phone": "phone number",
  "location": "City, State",
  "linkedin": "linkedin URL or empty string",
  "github": "github URL or empty string",
  "summary": "2-3 sentence professional summary based on their experience",
  "years_of_experience": <integer>,
  "current_title": "most recent job title",
  "target_titles": ["title1", "title2", "title3"],
  "skills": {{
    "languages": ["lang1", "lang2"],
    "frameworks": ["fw1", "fw2"],
    "infrastructure": ["infra1", "infra2"],
    "databases": ["db1", "db2"],
    "other": ["skill1", "skill2"]
  }},
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "start": "YYYY-MM",
      "end": "YYYY-MM or present",
      "bullets": ["achievement 1", "achievement 2"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree Name",
      "school": "School Name",
      "year": 2024
    }}
  ]
}}"""

    raw = chat_completion(system=system, user_message=user_msg, max_tokens=2000)

    # Extract JSON from response
    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not json_match:
        raise ValueError("AI could not parse the resume into structured data")

    parsed = json.loads(json_match.group())

    # Ensure all required fields exist with defaults
    parsed.setdefault("name", "")
    parsed.setdefault("email", "")
    parsed.setdefault("phone", "")
    parsed.setdefault("location", "")
    parsed.setdefault("linkedin", "")
    parsed.setdefault("github", "")
    parsed.setdefault("summary", "")
    parsed.setdefault("years_of_experience", 0)
    parsed.setdefault("current_title", "")
    parsed.setdefault("target_titles", [])
    parsed.setdefault("skills", {"languages": [], "frameworks": [], "infrastructure": [], "databases": [], "other": []})
    parsed.setdefault("experience", [])
    parsed.setdefault("education", [])
    parsed.setdefault("raw_resume_text", "")

    # Ensure skills has all sub-keys
    for key in ("languages", "frameworks", "infrastructure", "databases", "other"):
        parsed["skills"].setdefault(key, [])

    return parsed


@app.route("/api/profile/full", methods=["POST"])
def api_save_profile_full():
    """Save candidate profile and preferences from the setup form."""
    data = request.get_json()
    profile = data.get("profile")
    preferences = data.get("preferences")

    if not profile:
        return jsonify({"error": "Profile data is required"}), 400

    try:
        validate_profile(profile)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    save_profile(profile)
    if preferences:
        save_preferences(preferences)

    init_db()
    return jsonify({"status": "ok"})


# ---------------------------------------------------------------------------
# API — Resume Builder Files
# ---------------------------------------------------------------------------

_EXPERIENCE_DIR = PROJECT_ROOT / "data" / "experience"
_PROJECTS_DIR = PROJECT_ROOT / "data" / "projects"
_TEMPLATE_DIR = PROJECT_ROOT / "data" / "resume_template"


@app.route("/api/resume-builder/files")
def api_resume_builder_files():
    """Return all resume builder input files (experience, projects, template)."""
    files = {}

    for name, path in [
        ("experience_current", _EXPERIENCE_DIR / "current.md"),
        ("experience_previous", _EXPERIENCE_DIR / "previous.md"),
        ("projects", _PROJECTS_DIR / "projects.md"),
        ("template", _TEMPLATE_DIR / "template.tex"),
    ]:
        files[name] = path.read_text() if path.exists() else ""

    # Check readiness
    from pipeline.resume_builder_bridge import check_builder_ready
    status = check_builder_ready()

    return jsonify({"files": files, "status": status})


@app.route("/api/resume-builder/files", methods=["POST"])
def api_save_resume_builder_files():
    """Save resume builder input files."""
    data = request.get_json()
    files = data.get("files", {})

    _EXPERIENCE_DIR.mkdir(parents=True, exist_ok=True)
    _PROJECTS_DIR.mkdir(parents=True, exist_ok=True)
    _TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

    saved = []
    for name, path in [
        ("experience_current", _EXPERIENCE_DIR / "current.md"),
        ("experience_previous", _EXPERIENCE_DIR / "previous.md"),
        ("projects", _PROJECTS_DIR / "projects.md"),
        ("template", _TEMPLATE_DIR / "template.tex"),
    ]:
        if name in files:
            path.write_text(files[name])
            saved.append(name)

    return jsonify({"status": "ok", "saved": saved})


@app.route("/api/resume-builder/upload/<file_type>", methods=["POST"])
def api_resume_builder_upload(file_type):
    """Upload a file (PDF/DOCX/TXT) for experience or projects.
    Extracts text and saves as markdown."""
    import os as _os
    import tempfile as _tempfile

    valid_types = {
        "experience_current": _EXPERIENCE_DIR / "current.md",
        "experience_previous": _EXPERIENCE_DIR / "previous.md",
        "projects": _PROJECTS_DIR / "projects.md",
    }

    if file_type not in valid_types:
        return jsonify({"error": f"Invalid file type. Use: {', '.join(valid_types)}"}), 400

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = (file.filename or "").lower()

    if not any(filename.endswith(ext) for ext in (".pdf", ".docx", ".doc", ".txt", ".md")):
        return jsonify({"error": "Supported formats: PDF, DOCX, TXT, MD"}), 400

    suffix = "." + filename.rsplit(".", 1)[-1]
    with _tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        text = _extract_resume_text(tmp_path, suffix)
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "Could not extract text from file"}), 400

        # Save as markdown
        dest = valid_types[file_type]
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_text(text)

        return jsonify({"text": text, "chars": len(text), "saved_to": str(dest)})
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        _os.unlink(tmp_path)


@app.route("/api/resume-builder/status")
def api_resume_builder_status():
    """Check if the 3-stage resume builder is ready to use."""
    from pipeline.resume_builder_bridge import check_builder_ready
    return jsonify(check_builder_ready())


# ---------------------------------------------------------------------------
# API — Scrape
# ---------------------------------------------------------------------------

_scrape_progress = {
    "current": 0, "total": 4, "status": "idle", "message": "", "stats": None
}


@app.route("/api/scrape", methods=["POST"])
def api_scrape():
    """Trigger job scraping in a background thread."""
    if _scrape_progress["status"] == "running":
        return jsonify({"error": "Scrape already in progress"}), 409

    _scrape_progress["current"] = 0
    _scrape_progress["total"] = 4
    _scrape_progress["status"] = "running"
    _scrape_progress["message"] = "Starting scrapers..."
    _scrape_progress["stats"] = None

    def run_scrape():
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from pipeline.normalizer import normalize_and_insert
        init_db()
        prefs = load_preferences()

        # Run all 3 scrapers in parallel
        _scrape_progress["current"] = 1
        _scrape_progress["total"] = 3
        _scrape_progress["message"] = "Fetching from all sources in parallel..."

        all_jobs = []
        sources_done = 0

        def fetch_ats():
            from scrapers.ats_scraper import fetch_all_ats_jobs
            return "ATS", fetch_all_ats_jobs(prefs)

        def fetch_hiringcafe():
            from scrapers.hiringcafe_scraper import fetch_hiringcafe_jobs
            return "Hiring Cafe", fetch_hiringcafe_jobs(prefs)

        def fetch_jobspy():
            from scrapers.jobspy_scraper import scrape_major_boards
            return "JobSpy", scrape_major_boards(prefs)

        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = [
                executor.submit(fetch_ats),
                executor.submit(fetch_hiringcafe),
                executor.submit(fetch_jobspy),
            ]

            for future in as_completed(futures):
                try:
                    source_name, jobs = future.result()
                    all_jobs.extend(jobs)
                    sources_done += 1
                    _scrape_progress["current"] = sources_done
                    _scrape_progress["message"] = f"{source_name}: {len(jobs)} jobs fetched ({sources_done}/3 sources done)"
                except Exception as e:
                    sources_done += 1
                    _scrape_progress["current"] = sources_done
                    _scrape_progress["message"] = f"Source error: {e}"

        # Normalize and insert
        _scrape_progress["current"] = 3
        _scrape_progress["message"] = f"Normalizing {len(all_jobs)} jobs..."
        try:
            stats = normalize_and_insert(all_jobs)
            _scrape_progress["stats"] = stats
            _scrape_progress["message"] = (
                f"Done: {stats['new']} new jobs added, "
                f"{stats['duplicates']} duplicates skipped"
            )
        except Exception as e:
            _scrape_progress["message"] = f"Normalize error: {e}"

        _scrape_progress["status"] = "done"

    thread = threading.Thread(target=run_scrape, daemon=True)
    thread.start()
    return jsonify({"status": "started"})


@app.route("/api/scrape/progress")
def api_scrape_progress():
    """SSE endpoint for scrape progress."""
    def stream():
        import time
        while True:
            data = json.dumps(_scrape_progress)
            yield f"data: {data}\n\n"
            if _scrape_progress["status"] in ("done", "idle"):
                break
            time.sleep(0.5)
    return Response(stream(), mimetype="text/event-stream")


@app.route("/api/tasks/status")
def api_tasks_status():
    """Return the running/idle status of all background tasks.
    Frontend polls this on page load to re-attach progress bars."""
    return jsonify({
        "scrape": _scrape_progress["status"],
        "generate": _generation_progress["status"],
        "resume": _resume_progress["status"],
        "autofill": _autofill_progress["status"],
    })


# ---------------------------------------------------------------------------
# API — Score
# ---------------------------------------------------------------------------

_score_progress = {
    "current": 0, "total": 0, "status": "idle", "message": "", "stats": None
}


@app.route("/api/score", methods=["POST"])
def api_score():
    """Trigger job scoring in a background thread."""
    if _score_progress["status"] == "running":
        return jsonify({"error": "Scoring already in progress"}), 409

    _score_progress["current"] = 0
    _score_progress["total"] = 2
    _score_progress["status"] = "running"
    _score_progress["message"] = "Starting scorer..."
    _score_progress["stats"] = None

    def run_score():
        from pipeline.scorer import score_all_new_jobs
        init_db()
        profile = load_profile()
        prefs = load_preferences()

        # Step 1: Hard filters
        _score_progress["current"] = 1
        _score_progress["message"] = "Running hard filters and AI scoring..."

        try:
            stats = score_all_new_jobs(profile, prefs)
            _score_progress["current"] = 2
            _score_progress["stats"] = stats
            _score_progress["message"] = (
                f"Done: {stats.get('queued', 0)} queued, "
                f"{stats.get('scored', 0)} scored, "
                f"{stats.get('filtered_out', 0)} filtered out, "
                f"{stats.get('errors', 0)} errors"
            )
        except Exception as e:
            _score_progress["message"] = f"Scoring error: {e}"

        _score_progress["status"] = "done"

    thread = threading.Thread(target=run_score, daemon=True)
    thread.start()
    return jsonify({"status": "started"})


@app.route("/api/score/progress")
def api_score_progress():
    """SSE endpoint for score progress."""
    def stream():
        import time
        while True:
            data = json.dumps(_score_progress)
            yield f"data: {data}\n\n"
            if _score_progress["status"] in ("done", "idle"):
                break
            time.sleep(0.5)
    return Response(stream(), mimetype="text/event-stream")


# ---------------------------------------------------------------------------
# API — Job listing
# ---------------------------------------------------------------------------

def _job_to_dict(job: dict) -> dict:
    """Convert a job row to a JSON-safe dict for the frontend."""
    notes = {}
    if job.get("notes"):
        try:
            notes = json.loads(job["notes"])
        except (json.JSONDecodeError, TypeError):
            pass

    bullets = []
    if job.get("resume_bullets"):
        try:
            bullets = json.loads(job["resume_bullets"])
        except (json.JSONDecodeError, TypeError):
            pass

    resume_pdf_path = str(PROJECT_ROOT / "data" / "resumes" / f"{job['id']}.pdf")

    return {
        "id": job["id"],
        "title": job.get("title", ""),
        "company": job.get("company", ""),
        "location": job.get("location", ""),
        "url": job.get("url", ""),
        "source": job.get("source", ""),
        "score": job.get("score"),
        "score_reason": job.get("score_reason", ""),
        "status": job.get("status", ""),
        "has_cover_letter": bool(job.get("cover_letter")),
        "cover_letter": job.get("cover_letter", ""),
        "resume_bullets": bullets,
        "strengths": notes.get("strengths", []),
        "missing": notes.get("missing", []),
        "keywords": notes.get("keywords", []),
        "description": job.get("description", ""),
        "salary_min": job.get("salary_min"),
        "salary_max": job.get("salary_max"),
        "date_posted": job.get("date_posted", ""),
        "date_scraped": job.get("date_scraped", ""),
        "has_resume": os.path.exists(resume_pdf_path),
    }


@app.route("/api/jobs")
def api_jobs():
    """Fetch jobs by tab: queued, ready, submitted, skipped, new, scored, contract."""
    tab = request.args.get("tab", "queued")

    if tab == "queued":
        jobs = get_queued_without_cl()
    elif tab == "ready":
        jobs = get_queued_with_cl()
    elif tab == "submitted":
        jobs = get_jobs_by_status("submitted")
    elif tab == "skipped":
        jobs = get_jobs_by_status("skipped")
    elif tab == "new":
        jobs = get_jobs_by_status("new")
    elif tab == "scored":
        jobs = get_jobs_by_status("scored")
    elif tab == "contract":
        jobs = get_contract_jobs()
    else:
        jobs = get_jobs_by_status("queued")

    return jsonify([_job_to_dict(j) for j in jobs])


@app.route("/api/job/<job_id>")
def api_job_detail(job_id):
    """Get full details for a single job."""
    job = get_job_by_id(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(_job_to_dict(job))


@app.route("/api/stats")
def api_stats():
    """Get job counts by status."""
    counts = count_by_status()
    queued_no_cl = len(get_queued_without_cl())
    queued_with_cl = len(get_queued_with_cl())
    contract_count = len(get_contract_jobs())
    return jsonify({
        "queued": queued_no_cl,
        "ready": queued_with_cl,
        "submitted": counts.get("submitted", 0),
        "skipped": counts.get("skipped", 0),
        "scored": counts.get("scored", 0),
        "new": counts.get("new", 0),
        "contract": contract_count,
    })


# ---------------------------------------------------------------------------
# API — Actions
# ---------------------------------------------------------------------------

# Track generation progress for SSE
_generation_progress = {"current": 0, "total": 0, "status": "idle", "message": ""}


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """Generate cover letters + bullets for selected job IDs. Runs in background thread."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    if _generation_progress["status"] == "running":
        return jsonify({"error": "Generation already in progress"}), 409

    _generation_progress["current"] = 0
    _generation_progress["total"] = len(job_ids)
    _generation_progress["status"] = "running"
    _generation_progress["message"] = "Starting..."

    def run_generation():
        from pipeline.generator import generate_cover_letter, generate_resume_bullets
        profile = load_profile()
        generated = 0
        errors = 0

        for i, job_id in enumerate(job_ids):
            job = get_job_by_id(job_id)
            if not job:
                _generation_progress["current"] = i + 1
                continue

            _generation_progress["message"] = f"Cover letter: {job['title']} at {job['company']}"

            try:
                cl = generate_cover_letter(job, profile)
                bullets = generate_resume_bullets(job, profile)
                update_job(job_id, cover_letter=cl, resume_bullets=json.dumps(bullets))
                generated += 1
            except Exception as e:
                errors += 1
                _generation_progress["message"] = f"Error: {e}"

            _generation_progress["current"] = i + 1

        _generation_progress["status"] = "done"
        _generation_progress["message"] = f"Done: {generated} cover letters, {errors} errors"

    thread = threading.Thread(target=run_generation, daemon=True)
    thread.start()

    return jsonify({"status": "started", "total": len(job_ids)})


# Track resume generation progress
_resume_progress = {"current": 0, "total": 0, "status": "idle", "message": ""}


@app.route("/api/generate-resumes", methods=["POST"])
def api_generate_resumes():
    """Generate tailored resume PDFs using the 3-stage Claude CLI pipeline.
    Falls back to the built-in generator if the builder inputs aren't set up."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    if _resume_progress["status"] == "running":
        return jsonify({"error": "Resume generation already in progress"}), 409

    _resume_progress["current"] = 0
    _resume_progress["total"] = len(job_ids)
    _resume_progress["status"] = "running"
    _resume_progress["message"] = "Starting..."

    def run_resume_generation():
        from pipeline.resume_builder_bridge import generate_resume, check_builder_ready

        builder_status = check_builder_ready()
        use_3stage = builder_status["ready"]

        if use_3stage:
            _resume_progress["message"] = "Using 3-stage Claude CLI resume builder"
        else:
            _resume_progress["message"] = (
                "3-stage builder not ready — using built-in generator. "
                "Set up experience files in Settings to enable it."
            )

        generated = 0
        errors = 0

        for i, job_id in enumerate(job_ids):
            job = get_job_by_id(job_id)
            if not job:
                _resume_progress["current"] = i + 1
                continue

            _resume_progress["message"] = f"Resume {i+1}/{len(job_ids)}: {job['title']} at {job['company']}"

            try:
                if use_3stage:
                    def progress_cb(msg):
                        _resume_progress["message"] = f"[{i+1}/{len(job_ids)}] {msg}"

                    result = generate_resume(job, progress_callback=progress_cb)
                    if result["success"]:
                        generated += 1
                    else:
                        errors += 1
                        _resume_progress["message"] = f"Error: {result['error']}"
                else:
                    from pipeline.resume_generator import generate_tailored_resume
                    profile = load_profile()
                    generate_tailored_resume(job, profile)
                    generated += 1

            except Exception as e:
                errors += 1
                _resume_progress["message"] = f"Error: {e}"

            _resume_progress["current"] = i + 1

        method = "3-stage Claude CLI" if use_3stage else "built-in"
        _resume_progress["status"] = "done"
        _resume_progress["message"] = f"Done ({method}): {generated} resumes, {errors} errors"

    thread = threading.Thread(target=run_resume_generation, daemon=True)
    thread.start()

    return jsonify({"status": "started", "total": len(job_ids)})


@app.route("/api/generate-resumes/progress")
def api_resume_progress():
    """SSE endpoint for resume generation progress."""
    def stream():
        import time
        while True:
            data = json.dumps(_resume_progress)
            yield f"data: {data}\n\n"
            if _resume_progress["status"] in ("done", "idle"):
                break
            time.sleep(0.5)
    return Response(stream(), mimetype="text/event-stream")


@app.route("/api/generate/progress")
def api_generate_progress():
    """SSE endpoint for generation progress."""
    def stream():
        import time
        while True:
            data = json.dumps(_generation_progress)
            yield f"data: {data}\n\n"
            if _generation_progress["status"] in ("done", "idle"):
                break
            time.sleep(0.5)
    return Response(stream(), mimetype="text/event-stream")


@app.route("/api/cover-letters/download-all", methods=["POST"])
def api_download_all_cover_letters():
    """Download multiple cover letters as a single ZIP file."""
    import io
    import zipfile

    data = request.get_json()
    job_ids = data.get("job_ids", [])
    fmt = data.get("format", "pdf").lower()

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    profile = load_profile()
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for job_id in job_ids:
            job = get_job_by_id(job_id)
            if not job or not job.get("cover_letter"):
                continue

            cover_letter = job["cover_letter"]
            company = job.get("company", "Company").replace(" ", "_").replace("/", "-")
            title = job.get("title", "Role").replace(" ", "_").replace("/", "-")
            base_name = f"Cover_Letter_{company}_{title}"

            if fmt == "docx":
                resp = _generate_docx(cover_letter, job, profile, base_name)
                zf.writestr(f"{base_name}.docx", resp.get_data())
            elif fmt == "txt":
                zf.writestr(f"{base_name}.txt", cover_letter.encode("utf-8"))
            else:
                resp = _generate_pdf(cover_letter, job, profile, base_name)
                zf.writestr(f"{base_name}.pdf", resp.get_data())

    buf.seek(0)

    ext_label = {"pdf": "PDF", "docx": "Word", "txt": "Text"}.get(fmt, "PDF")
    return Response(
        buf.getvalue(),
        mimetype="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="Cover_Letters_{ext_label}_{len(job_ids)}_jobs.zip"'
        },
    )


@app.route("/api/job/<job_id>/cover-letter/download")
def api_download_cover_letter(job_id):
    """Download a job's cover letter as PDF, DOCX, or TXT."""
    job = get_job_by_id(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    cover_letter = job.get("cover_letter") or ""
    if not cover_letter:
        return jsonify({"error": "No cover letter generated"}), 404

    fmt = request.args.get("format", "pdf").lower()
    company = job.get("company", "Company").replace(" ", "_")
    title = job.get("title", "Role").replace(" ", "_")
    base_name = f"Cover_Letter_{company}_{title}"

    profile = load_profile()

    if fmt == "docx":
        return _generate_docx(cover_letter, job, profile, base_name)
    elif fmt == "txt":
        return Response(
            cover_letter,
            mimetype="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.txt"'},
        )
    else:
        return _generate_pdf(cover_letter, job, profile, base_name)


@app.route("/api/job/<job_id>/resume/download")
def api_download_resume(job_id):
    """Download a job's tailored resume PDF. Returns 404 if not generated yet."""
    job = get_job_by_id(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    resume_path = PROJECT_ROOT / "data" / "resumes" / f"{job_id}.pdf"
    if not resume_path.exists():
        return jsonify({"error": "Tailored resume not generated yet"}), 404

    company = job.get("company", "Company").replace(" ", "_").replace("/", "-")
    title = job.get("title", "Role").replace(" ", "_").replace("/", "-")
    download_name = f"Resume_{company}_{title}.pdf"

    return send_file(
        str(resume_path),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=download_name,
    )


@app.route("/api/resumes/download-all", methods=["POST"])
def api_download_all_resumes():
    """Download multiple tailored resume PDFs as a single ZIP."""
    import io
    import zipfile

    data = request.get_json()
    job_ids = data.get("job_ids", [])
    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    buf = io.BytesIO()
    count = 0
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for job_id in job_ids:
            resume_path = PROJECT_ROOT / "data" / "resumes" / f"{job_id}.pdf"
            if not resume_path.exists():
                continue
            job = get_job_by_id(job_id)
            company = (job.get("company", "Company") if job else "Company").replace(" ", "_").replace("/", "-")
            title = (job.get("title", "Role") if job else "Role").replace(" ", "_").replace("/", "-")
            zf.write(str(resume_path), f"Resume_{company}_{title}.pdf")
            count += 1

    if count == 0:
        return jsonify({"error": "No resumes generated yet"}), 404

    buf.seek(0)
    return Response(
        buf.getvalue(),
        mimetype="application/zip",
        headers={"Content-Disposition": f'attachment; filename="Tailored_Resumes_{count}_jobs.zip"'},
    )


def _register_pdf_fonts(pdf):
    """Register Arial TTF fonts for Unicode support across macOS, Linux, Windows.
    Falls back to Helvetica if no TTF found."""
    import os
    candidates = {
        "regular": [
            "/System/Library/Fonts/Supplemental/Arial.ttf",                     # macOS
            "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",                # Linux
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",  # Linux fallback
            "C:\\Windows\\Fonts\\arial.ttf",                                    # Windows
        ],
        "bold": [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
            "C:\\Windows\\Fonts\\arialbd.ttf",
        ],
        "italic": [
            "/System/Library/Fonts/Supplemental/Arial Italic.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/Arial_Italic.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
            "C:\\Windows\\Fonts\\ariali.ttf",
        ],
    }

    def find(paths):
        for p in paths:
            if os.path.isfile(p):
                return p
        return None

    regular = find(candidates["regular"])
    bold = find(candidates["bold"])

    if regular and bold:
        pdf.add_font("arial", "", regular, uni=True)
        pdf.add_font("arial", "B", bold, uni=True)
        italic = find(candidates["italic"])
        if italic:
            pdf.add_font("arial", "I", italic, uni=True)
        return "arial"

    return "Helvetica"


def _generate_pdf(cover_letter: str, job: dict, profile: dict, base_name: str):
    """Generate a professional cover letter PDF with Unicode support."""
    import io
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)

    font = _register_pdf_fonts(pdf)

    pdf.add_page()

    # Header — candidate name
    pdf.set_font(font, "B", 18)
    pdf.cell(0, 10, profile.get("name", ""), new_x="LMARGIN", new_y="NEXT")

    # Contact line
    pdf.set_font(font, "", 9)
    pdf.set_text_color(100, 100, 100)
    contact_parts = [
        profile.get("email", ""),
        profile.get("phone", ""),
        profile.get("location", ""),
    ]
    pdf.cell(0, 5, "  |  ".join(p for p in contact_parts if p), new_x="LMARGIN", new_y="NEXT")

    # LinkedIn / GitHub line
    links = [profile.get("linkedin", ""), profile.get("github", "")]
    links_str = "  |  ".join(l for l in links if l)
    if links_str:
        pdf.cell(0, 5, links_str, new_x="LMARGIN", new_y="NEXT")

    # Divider line
    pdf.set_draw_color(200, 200, 200)
    pdf.ln(4)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Date and target
    from datetime import datetime
    pdf.set_text_color(100, 100, 100)
    pdf.set_font(font, "", 10)
    pdf.cell(0, 5, datetime.now().strftime("%B %d, %Y"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font(font, "B", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 5, f"{job.get('title', '')} at {job.get('company', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Cover letter body
    pdf.set_font(font, "", 11)
    pdf.set_text_color(30, 30, 30)
    for paragraph in cover_letter.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            pdf.multi_cell(0, 6, paragraph)
            pdf.ln(4)

    # Sign off
    pdf.ln(4)
    pdf.set_font(font, "", 11)
    pdf.cell(0, 6, "Best regards,", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font, "B", 11)
    pdf.cell(0, 6, profile.get("name", ""), new_x="LMARGIN", new_y="NEXT")

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    return Response(
        buf.getvalue(),
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base_name}.pdf"'},
    )


def _generate_docx(cover_letter: str, job: dict, profile: dict, base_name: str):
    """Generate a professional cover letter DOCX."""
    import io
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header — name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_para.space_after = Pt(2)

    # Contact info
    contact_parts = [
        profile.get("email", ""),
        profile.get("phone", ""),
        profile.get("location", ""),
    ]
    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run("  |  ".join(p for p in contact_parts if p))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(100, 100, 100)
    contact_para.space_after = Pt(1)

    # Links
    links = [profile.get("linkedin", ""), profile.get("github", "")]
    links_str = "  |  ".join(l for l in links if l)
    if links_str:
        links_para = doc.add_paragraph()
        links_run = links_para.add_run(links_str)
        links_run.font.size = Pt(9)
        links_run.font.color.rgb = RGBColor(100, 100, 100)
        links_para.space_after = Pt(6)

    # Divider
    divider = doc.add_paragraph()
    divider_run = divider.add_run("_" * 80)
    divider_run.font.size = Pt(6)
    divider_run.font.color.rgb = RGBColor(200, 200, 200)
    divider.space_after = Pt(8)

    # Date
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    date_para.space_after = Pt(2)

    # Target role
    role_para = doc.add_paragraph()
    role_run = role_para.add_run(f"{job.get('title', '')} at {job.get('company', '')}")
    role_run.bold = True
    role_run.font.size = Pt(10)
    role_para.space_after = Pt(12)

    # Body
    for paragraph in cover_letter.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            p = doc.add_paragraph(paragraph)
            p.space_after = Pt(8)

    # Sign-off
    doc.add_paragraph()
    regards = doc.add_paragraph("Best regards,")
    regards.space_after = Pt(2)
    sign = doc.add_paragraph()
    sign_run = sign.add_run(profile.get("name", ""))
    sign_run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        buf.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{base_name}.docx"'},
    )


@app.route("/api/skip", methods=["POST"])
def api_skip():
    """Skip selected jobs."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])
    for job_id in job_ids:
        update_job(job_id, status="skipped", score_reason="[MANUAL] Skipped via dashboard")
    return jsonify({"status": "ok", "skipped": len(job_ids)})


@app.route("/api/job/<job_id>/mark-applied", methods=["POST"])
def api_mark_applied(job_id):
    """Quick mark a single job as submitted (applied externally)."""
    from datetime import datetime, timezone
    job = get_job_by_id(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    update_job(job_id, status="submitted", date_submitted=datetime.now(timezone.utc).isoformat())
    from utils.db import insert_application
    insert_application(job_id, "Applied externally (marked via dashboard)")
    return jsonify({"status": "ok"})


@app.route("/api/job/<job_id>/skip", methods=["POST"])
def api_skip_single(job_id):
    """Quick skip/dismiss a single job."""
    job = get_job_by_id(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    update_job(job_id, status="skipped", score_reason="[MANUAL] Dismissed via dashboard")
    return jsonify({"status": "ok"})


@app.route("/api/job/<job_id>/cover-letter", methods=["PUT"])
def api_update_cover_letter(job_id):
    """Update a job's cover letter."""
    data = request.get_json()
    cover_letter = data.get("cover_letter", "")
    update_job(job_id, cover_letter=cover_letter)
    return jsonify({"status": "ok"})


@app.route("/api/profile")
def api_profile():
    """Return the candidate profile fields that will be used for applications."""
    profile = load_profile()
    name_parts = profile["name"].split(" ", 1)
    return jsonify({
        "first_name": name_parts[0],
        "last_name": name_parts[1] if len(name_parts) > 1 else "",
        "email": profile.get("email", ""),
        "phone": profile.get("phone", ""),
        "location": profile.get("location", ""),
        "linkedin": profile.get("linkedin", ""),
        "github": profile.get("github", ""),
    })


@app.route("/api/submit/preview", methods=["POST"])
def api_submit_preview():
    """Return preview data for selected jobs before submission."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    profile = load_profile()
    name_parts = profile["name"].split(" ", 1)

    profile_fields = {
        "first_name": name_parts[0],
        "last_name": name_parts[1] if len(name_parts) > 1 else "",
        "email": profile.get("email", ""),
        "phone": profile.get("phone", ""),
        "location": profile.get("location", ""),
        "linkedin": profile.get("linkedin", ""),
        "github": profile.get("github", ""),
    }

    jobs_preview = []
    for job_id in job_ids:
        job = get_job_by_id(job_id)
        if not job:
            continue
        jobs_preview.append({
            "id": job["id"],
            "title": job.get("title", ""),
            "company": job.get("company", ""),
            "url": job.get("url", ""),
            "source": job.get("source", ""),
            "has_cover_letter": bool(job.get("cover_letter")),
            "cover_letter_preview": (job.get("cover_letter") or "")[:200] + ("..." if len(job.get("cover_letter") or "") > 200 else ""),
        })

    return jsonify({"profile": profile_fields, "jobs": jobs_preview})


@app.route("/api/submit", methods=["POST"])
def api_submit():
    """Mark selected jobs as submitted and return their URLs for the browser to open."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    results = []
    apply_urls = []
    for job_id in job_ids:
        job = get_job_by_id(job_id)
        if not job:
            results.append({"id": job_id, "status": "skipped", "reason": "Job not found"})
            continue

        update_job(job_id, status="submitted")
        from utils.db import insert_application
        insert_application(job_id, "Submitted via web dashboard")
        apply_urls.append(f"/apply/{job_id}")
        results.append({
            "id": job_id,
            "status": "submitted",
            "url": job["url"],
            "title": job.get("title", ""),
            "company": job.get("company", ""),
        })

    return jsonify({"results": results, "urls": apply_urls})


# ---------------------------------------------------------------------------
# API — Auto-fill (Playwright)
# ---------------------------------------------------------------------------

_autofill_progress = {"current": 0, "total": 0, "status": "idle", "message": "", "results": []}


@app.route("/api/autofill", methods=["POST"])
def api_autofill():
    """Launch Playwright to auto-fill application forms for selected jobs."""
    data = request.get_json()
    job_ids = data.get("job_ids", [])

    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    if _autofill_progress["status"] == "running":
        return jsonify({"error": "Auto-fill already in progress"}), 409

    _autofill_progress["current"] = 0
    _autofill_progress["total"] = len(job_ids)
    _autofill_progress["status"] = "running"
    _autofill_progress["message"] = "Launching browser..."
    _autofill_progress["results"] = []

    def progress_callback(current, total, job, result):
        _autofill_progress["current"] = current
        if job:
            status = result.get("status", "unknown")
            filled_count = len(result.get("filled", {}))
            fields = ", ".join(result.get("filled", {}).keys()) if result.get("filled") else "none"
            _autofill_progress["message"] = f"{job.get('title', '')} at {job.get('company', '')} — {filled_count} fields filled"
            _autofill_progress["results"].append({
                "title": job.get("title", ""),
                "company": job.get("company", ""),
                "ats_type": result.get("ats_type", ""),
                "status": status,
                "filled_fields": list(result.get("filled", {}).keys()),
                "filled_details": result.get("filled", {}),
            })
        else:
            _autofill_progress["message"] = "Job not found, skipping..."

    def run_autofill():
        from pipeline.submitter import autofill_jobs_sync
        try:
            autofill_jobs_sync(job_ids, progress_callback)
        except Exception as e:
            _autofill_progress["message"] = f"Error: {e}"
        finally:
            _autofill_progress["status"] = "done"
            _autofill_progress["message"] = f"Done — {len(_autofill_progress['results'])} jobs processed"

            # Mark jobs as submitted
            for job_id in job_ids:
                job = get_job_by_id(job_id)
                if job:
                    update_job(job_id, status="submitted")
                    from utils.db import insert_application
                    insert_application(job_id, "Auto-filled via dashboard")

    thread = threading.Thread(target=run_autofill, daemon=True)
    thread.start()

    return jsonify({"status": "started", "total": len(job_ids)})


@app.route("/api/autofill/progress")
def api_autofill_progress():
    """SSE endpoint for auto-fill progress."""
    def stream():
        import time
        while True:
            data = json.dumps(_autofill_progress)
            yield f"data: {data}\n\n"
            if _autofill_progress["status"] in ("done", "idle"):
                break
            time.sleep(0.5)
    return Response(stream(), mimetype="text/event-stream")


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------

def run_dashboard(port=5050):
    """Start the dashboard server and open browser."""
    import webbrowser
    init_db()
    webbrowser.open(f"http://localhost:{port}")
    app.run(host="127.0.0.1", port=port, debug=False)
